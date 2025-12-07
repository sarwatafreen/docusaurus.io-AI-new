import os
import re
from typing import List, Dict, Any
import markdown
from pathlib import Path
import logging

# Try to import optional dependencies
try:
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PyPDF2 not available. PDF support will be disabled.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx not available. DOCX support will be disabled.")

try:
    import trafilatura
    TRAFILATURA_SUPPORT = True
except ImportError:
    TRAFILATURA_SUPPORT = False
    logging.warning("trafilatura not available. Web content extraction support will be disabled.")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    A class to process various document formats for the RAG system.
    """
    
    def __init__(self):
        self.supported_formats = ['.md', '.txt']
        
        if PDF_SUPPORT:
            self.supported_formats.append('.pdf')
        
        if DOCX_SUPPORT:
            self.supported_formats.append('.docx')
    
    def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a single file and return a list of document chunks with metadata.
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}. Supported formats: {self.supported_formats}")
        
        if extension == '.pdf' and PDF_SUPPORT:
            return self._process_pdf(path)
        elif extension == '.docx' and DOCX_SUPPORT:
            return self._process_docx(path)
        elif extension == '.md':
            return self._process_markdown(path)
        elif extension == '.txt':
            return self._process_text(path)
        else:
            raise ValueError(f"Could not process {extension} file due to missing dependencies")
    
    def _process_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process a PDF file and return its content as chunks.
        """
        if not PDF_SUPPORT:
            raise RuntimeError("PyPDF2 is required to process PDF files")
        
        chunks = []
        try:
            with open(file_path, 'rb') as pdf_file:
                reader = pypdf.PdfReader(pdf_file)
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    
                    # Split into chunks of reasonable size (e.g., ~500 words)
                    text_chunks = self._chunk_text(text)
                    
                    for i, chunk in enumerate(text_chunks):
                        chunks.append({
                            'content': chunk,
                            'metadata': {
                                'source': str(file_path),
                                'page': page_num + 1,
                                'chunk_id': f"{file_path.name}_page{page_num+1}_chunk{i+1}",
                                'type': 'pdf'
                            }
                        })
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process a DOCX file and return its content as chunks.
        """
        if not DOCX_SUPPORT:
            raise RuntimeError("python-docx is required to process DOCX files")
        
        chunks = []
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Also get text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            text = '\n'.join(full_text)
            
            # Split into chunks of reasonable size
            text_chunks = self._chunk_text(text)
            
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    'content': chunk,
                    'metadata': {
                        'source': str(file_path),
                        'chunk_id': f"{file_path.name}_chunk{i+1}",
                        'type': 'docx'
                    }
                })
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_markdown(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process a Markdown file and return its content as chunks.
        """
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Convert markdown to plain text for processing
            plain_text = markdown.markdown(content)
            # Remove HTML tags to get plain text
            plain_text = re.sub('<[^<]+?>', '', plain_text)
            
            # Split into chunks
            text_chunks = self._chunk_text(plain_text)
            
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    'content': chunk,
                    'metadata': {
                        'source': str(file_path),
                        'chunk_id': f"{file_path.name}_chunk{i+1}",
                        'type': 'markdown'
                    }
                })
        except Exception as e:
            logger.error(f"Error processing Markdown file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _process_text(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process a plain text file and return its content as chunks.
        """
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into chunks
            text_chunks = self._chunk_text(content)
            
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    'content': chunk,
                    'metadata': {
                        'source': str(file_path),
                        'chunk_id': f"{file_path.name}_chunk{i+1}",
                        'type': 'text'
                    }
                })
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _chunk_text(self, text: str, max_chunk_size: int = 1000) -> List[str]:
        """
        Split text into chunks of maximum specified size, trying to break at sentence boundaries.
        """
        # First, try to split by paragraphs
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If adding this paragraph would exceed the limit
            if len(current_chunk) + len(paragraph) > max_chunk_size:
                # If current chunk is not empty, save it
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                
                # If the paragraph itself is larger than the limit, split it further
                if len(paragraph) > max_chunk_size:
                    # Split the large paragraph into sentences
                    sentences = re.split(r'[.!?]+', paragraph)
                    temp_chunk = ""
                    
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if not sentence:
                            continue
                        
                        if len(temp_chunk) + len(sentence) > max_chunk_size:
                            if temp_chunk.strip():
                                chunks.append(temp_chunk.strip())
                                temp_chunk = sentence + ". "
                            else:
                                temp_chunk = sentence + ". "
                        else:
                            temp_chunk += sentence + ". "
                    
                    if temp_chunk.strip():
                        current_chunk = temp_chunk
                else:
                    current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[Dict[str, Any]]:
        """
        Process all supported files in a directory and return their content as chunks.
        """
        directory = Path(directory_path)
        all_chunks = []
        
        if recursive:
            files = directory.rglob('*')
        else:
            files = directory.glob('*')
        
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    file_chunks = self.process_file(str(file_path))
                    all_chunks.extend(file_chunks)
                    logger.info(f"Processed {file_path} - {len(file_chunks)} chunks")
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
        
        return all_chunks


# Example usage function
def main():
    processor = DocumentProcessor()
    # Example: Process a directory containing book content
    # chunks = processor.process_directory('./book_content')
    # print(f"Processed {len(chunks)} chunks")


if __name__ == "__main__":
    main()